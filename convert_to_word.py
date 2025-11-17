#!/usr/bin/env python3
"""
Script để convert Markdown báo cáo sang Word document
Yêu cầu: pip install python-docx markdown
"""

import sys
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Cần cài đặt python-docx:")
    print("pip install python-docx")
    sys.exit(1)

def create_word_document():
    doc = Document()
    
    # Set font for Vietnamese
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('BÁO CÁO ỨNG DỤNG COOKING GUIDE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Ứng dụng Công thức Nấu ăn trên Android')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Blank line
    
    # Table of Contents
    doc.add_heading('MỤC LỤC', 1)
    toc_items = [
        '1. Tổng quan dự án',
        '2. Kiến trúc hệ thống',
        '3. Sơ đồ Use Case',
        '4. Phân tích chức năng',
        '5. Công nghệ sử dụng',
        '6. Cấu trúc dự án',
        '7. Giao diện người dùng',
        '8. Kết luận'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Section 1
    doc.add_heading('1. TỔNG QUAN DỰ ÁN', 1)
    doc.add_heading('1.1. Giới thiệu', 2)
    doc.add_paragraph(
        'Cooking Guide là ứng dụng di động được phát triển bằng Flutter, cho phép người dùng:'
    )
    features = [
        'Quản lý và xem các công thức nấu ăn',
        'Tìm kiếm công thức theo tên hoặc mô tả',
        'Lọc công thức theo danh mục',
        'Xem chi tiết công thức với nguyên liệu và các bước thực hiện'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.2. Mục tiêu dự án', 2)
    goals = [
        'Cung cấp nền tảng dễ sử dụng để quản lý công thức nấu ăn',
        'Tích hợp Firebase để lưu trữ dữ liệu và xác thực người dùng',
        'Giao diện hiện đại, thân thiện với người dùng',
        'Hỗ trợ tìm kiếm và lọc công thức hiệu quả'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_heading('1.3. Phạm vi ứng dụng', 2)
    doc.add_paragraph('Platform: Android (có thể mở rộng sang iOS, Web)')
    doc.add_paragraph('Đối tượng sử dụng: Người dùng cá nhân quan tâm đến nấu ăn')
    doc.add_paragraph('Chức năng chính: Quản lý, tìm kiếm và xem công thức nấu ăn')
    
    doc.add_page_break()
    
    # Section 2
    doc.add_heading('2. KIẾN TRÚC HỆ THỐNG', 1)
    doc.add_heading('2.1. Kiến trúc tổng quan', 2)
    
    # Architecture diagram as text
    arch_text = """
┌─────────────────────────────────────────────────────────┐
│                    FLUTTER APP                          │
├─────────────────────────────────────────────────────────┤
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐ │
│  │   Screens    │  │   Widgets    │  │   Services   │ │
│  │              │  │              │  │              │ │
│  │ - Login      │  │ - RecipeCard │  │ - Firebase   │ │
│  │ - Home       │  │              │  │   Service    │ │
│  │ - Detail     │  │              │  │              │ │
│  └──────────────┘  └──────────────┘  └──────────────┘ │
│         │                  │                  │        │
│         └──────────────────┴──────────────────┘        │
│                         │                               │
│                  ┌──────────────┐                      │
│                  │   Models     │                      │
│                  │ - Recipe      │                      │
│                  └──────────────┘                      │
└─────────────────────────│───────────────────────────────┘
                          │
                          ▼
┌─────────────────────────────────────────────────────────┐
│                    FIREBASE BACKEND                     │
├─────────────────────────────────────────────────────────┤
│  ┌──────────────┐          ┌──────────────┐           │
│  │   Firebase   │          │  Cloud       │           │
│  │   Auth       │          │  Firestore   │           │
│  │              │          │              │           │
│  │ - Login      │          │ - Recipes    │           │
│  │ - Signup     │          │   Collection │           │
│  └──────────────┘          └──────────────┘           │
└─────────────────────────────────────────────────────────┘
"""
    doc.add_paragraph(arch_text)
    
    doc.add_heading('2.2. Luồng dữ liệu', 2)
    flows = [
        'Xác thực người dùng: Login Screen → Firebase Auth → Home Screen',
        'Tải công thức: Home Screen → Firebase Service → Firestore → Hiển thị',
        'Xem chi tiết: Home Screen → Recipe Detail Screen'
    ]
    for flow in flows:
        doc.add_paragraph(flow, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 3 - Use Case Diagram
    doc.add_heading('3. SƠ ĐỒ USE CASE', 1)
    doc.add_heading('3.1. Sơ đồ Use Case tổng quan', 2)
    
    use_case_diagram = """
                    ┌─────────────────┐
                    │  👤 Người dùng  │
                    └────────┬────────┘
                             │
        ┌────────────────────┼────────────────────┐
        │                    │                    │
        ▼                    ▼                    ▼
   ┌─────────┐         ┌─────────┐         ┌─────────┐
   │  UC1    │         │  UC2    │         │  UC3    │
   │ Đăng ký │         │Đăng nhập│         │Xem danh │
   │         │         │         │         │sách CT  │
   └─────────┘         └────┬────┘         └────┬────┘
                             │                   │
                             │                   │
        ┌────────────────────┼───────────────────┼────────────┐
        │                    │                   │            │
        ▼                    ▼                   ▼            ▼
   ┌─────────┐         ┌─────────┐         ┌─────────┐ ┌─────────┐
   │  UC4    │         │  UC5    │         │  UC6    │ │  UC7    │
   │Tìm kiếm │         │Lọc theo │         │Xem chi  │ │Đăng xuất│
   │         │         │danh mục │         │tiết     │ │         │
   └─────────┘         └─────────┘         └─────────┘ └─────────┘
"""
    doc.add_paragraph(use_case_diagram)
    
    doc.add_heading('3.2. Mô tả Use Case', 2)
    
    # UC1
    doc.add_heading('UC1: Đăng ký tài khoản', 3)
    doc.add_paragraph('Actor: Người dùng chưa có tài khoản')
    doc.add_paragraph('Mô tả: Người dùng tạo tài khoản mới bằng email và mật khẩu')
    doc.add_paragraph('Precondition: Ứng dụng đã được mở')
    doc.add_paragraph('Postcondition: Tài khoản được tạo, người dùng được đăng nhập tự động')
    doc.add_paragraph('Luồng chính:')
    uc1_steps = [
        'Người dùng nhập email và mật khẩu',
        'Hệ thống validate dữ liệu',
        'Gửi yêu cầu đăng ký đến Firebase Auth',
        'Tài khoản được tạo thành công',
        'Chuyển đến màn hình chính'
    ]
    for step in uc1_steps:
        doc.add_paragraph(step, style='List Number')
    
    # UC2
    doc.add_heading('UC2: Đăng nhập', 3)
    doc.add_paragraph('Actor: Người dùng đã có tài khoản')
    doc.add_paragraph('Mô tả: Người dùng đăng nhập vào hệ thống')
    doc.add_paragraph('Precondition: Người dùng đã có tài khoản')
    doc.add_paragraph('Postcondition: Người dùng được xác thực và chuyển đến màn hình chính')
    doc.add_paragraph('Luồng chính:')
    uc2_steps = [
        'Người dùng nhập email và mật khẩu',
        'Hệ thống validate dữ liệu',
        'Gửi yêu cầu đăng nhập đến Firebase Auth',
        'Xác thực thành công',
        'Chuyển đến màn hình chính'
    ]
    for step in uc2_steps:
        doc.add_paragraph(step, style='List Number')
    
    # UC3
    doc.add_heading('UC3: Xem danh sách công thức', 3)
    doc.add_paragraph('Actor: Người dùng đã đăng nhập')
    doc.add_paragraph('Mô tả: Hiển thị danh sách tất cả công thức nấu ăn')
    doc.add_paragraph('Precondition: Người dùng đã đăng nhập')
    doc.add_paragraph('Postcondition: Danh sách công thức được hiển thị')
    
    # UC4
    doc.add_heading('UC4: Tìm kiếm công thức', 3)
    doc.add_paragraph('Actor: Người dùng đã đăng nhập')
    doc.add_paragraph('Mô tả: Tìm kiếm công thức theo từ khóa')
    
    # UC5
    doc.add_heading('UC5: Lọc công thức theo danh mục', 3)
    doc.add_paragraph('Actor: Người dùng đã đăng nhập')
    doc.add_paragraph('Mô tả: Lọc công thức theo danh mục')
    
    # UC6
    doc.add_heading('UC6: Xem chi tiết công thức', 3)
    doc.add_paragraph('Actor: Người dùng đã đăng nhập')
    doc.add_paragraph('Mô tả: Xem thông tin chi tiết của một công thức')
    
    # UC7
    doc.add_heading('UC7: Đăng xuất', 3)
    doc.add_paragraph('Actor: Người dùng đã đăng nhập')
    doc.add_paragraph('Mô tả: Đăng xuất khỏi tài khoản')
    
    doc.add_page_break()
    
    # Section 4
    doc.add_heading('4. PHÂN TÍCH CHỨC NĂNG', 1)
    doc.add_heading('4.1. Màn hình đăng nhập/đăng ký', 2)
    doc.add_paragraph('Chức năng:')
    login_features = [
        'Đăng nhập với email và mật khẩu',
        'Đăng ký tài khoản mới',
        'Chuyển đổi giữa chế độ đăng nhập và đăng ký',
        'Validation form',
        'Hiển thị/ẩn mật khẩu',
        'Xử lý lỗi chi tiết'
    ]
    for feature in login_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('4.2. Màn hình chính', 2)
    doc.add_paragraph('Chức năng:')
    home_features = [
        'Hiển thị danh sách công thức',
        'Tìm kiếm công thức',
        'Lọc theo danh mục',
        'Pull-to-refresh',
        'Đăng xuất'
    ]
    for feature in home_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('4.3. Màn hình chi tiết công thức', 2)
    doc.add_paragraph('Chức năng:')
    detail_features = [
        'Hiển thị hình ảnh công thức',
        'Hiển thị mô tả',
        'Danh sách nguyên liệu',
        'Các bước thực hiện'
    ]
    for feature in detail_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 5
    doc.add_heading('5. CÔNG NGHỆ SỬ DỤNG', 1)
    doc.add_heading('5.1. Framework và ngôn ngữ', 2)
    doc.add_paragraph('Flutter: Framework đa nền tảng')
    doc.add_paragraph('Dart: Ngôn ngữ lập trình (version 3.9.2+)')
    doc.add_paragraph('Material Design 3: Design system')
    
    doc.add_heading('5.2. Backend và Database', 2)
    doc.add_paragraph('Firebase Core: 4.2.0')
    doc.add_paragraph('Firebase Authentication: 6.1.1')
    doc.add_paragraph('Cloud Firestore: 6.0.3')
    
    doc.add_heading('5.3. Công cụ phát triển', 2)
    doc.add_paragraph('Android Studio: IDE chính')
    doc.add_paragraph('Flutter SDK: 3.35.7')
    doc.add_paragraph('Gradle: 8.11.1')
    
    doc.add_page_break()
    
    # Section 6
    doc.add_heading('6. CẤU TRÚC DỰ ÁN', 1)
    structure_text = """
lib/
├── constants/
│   ├── app_colors.dart      # Định nghĩa màu sắc
│   └── app_styles.dart      # Định nghĩa styles
├── models/
│   └── recipe_model.dart    # Model công thức
├── screens/
│   ├── login_screen.dart    # Màn hình đăng nhập/đăng ký
│   ├── home_screen.dart     # Màn hình chính
│   └── recipe_detail_screen.dart # Màn hình chi tiết
├── services/
│   └── firebase_service.dart # Service tương tác Firebase
├── widgets/
│   └── recipe_card.dart     # Widget hiển thị card công thức
├── firebase_options.dart     # Cấu hình Firebase
└── main.dart                 # Entry point
"""
    doc.add_paragraph(structure_text)
    
    doc.add_page_break()
    
    # Section 7
    doc.add_heading('7. GIAO DIỆN NGƯỜI DÙNG', 1)
    doc.add_heading('7.1. Design System', 2)
    doc.add_paragraph('Màu sắc:')
    doc.add_paragraph('Primary: #FF6B35 (Cam)')
    doc.add_paragraph('Accent: #FFB84D (Vàng cam)')
    doc.add_paragraph('Background: #F8F9FA (Xám nhạt)')
    
    doc.add_heading('7.2. Các màn hình', 2)
    doc.add_paragraph('Login Screen: Gradient background, form validation')
    doc.add_paragraph('Home Screen: Search bar, category filters, recipe cards')
    doc.add_paragraph('Recipe Detail Screen: SliverAppBar, numbered lists')
    
    doc.add_page_break()
    
    # Section 8
    doc.add_heading('8. KẾT LUẬN', 1)
    doc.add_heading('8.1. Tóm tắt', 2)
    doc.add_paragraph('Ứng dụng Cooking Guide đã được phát triển thành công với các chức năng chính:')
    summary = [
        'Xác thực người dùng với Firebase Auth',
        'Quản lý và hiển thị công thức nấu ăn',
        'Tìm kiếm và lọc công thức',
        'Giao diện hiện đại, thân thiện'
    ]
    for item in summary:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('8.2. Điểm mạnh', 2)
    strengths = [
        'Kiến trúc rõ ràng, dễ bảo trì',
        'Giao diện đẹp, UX tốt',
        'Xử lý lỗi chi tiết',
        'Code được tổ chức tốt'
    ]
    for strength in strengths:
        doc.add_paragraph(strength, style='List Bullet')
    
    doc.add_heading('8.3. Hướng phát triển', 2)
    future = [
        'Thêm chức năng tạo/sửa/xóa công thức',
        'Thêm chức năng yêu thích',
        'Thêm đánh giá và bình luận',
        'Tối ưu hiệu năng và caching',
        'Hỗ trợ offline mode'
    ]
    for item in future:
        doc.add_paragraph(item, style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph(f'Ngày tạo báo cáo: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph('Phiên bản ứng dụng: 1.0.0+1')
    doc.add_paragraph('Framework: Flutter 3.35.7')
    doc.add_paragraph('Platform: Android')
    
    # Save document
    output_file = 'BAO_CAO_COOKING_GUIDE.docx'
    doc.save(output_file)
    print(f'✅ Đã tạo file Word: {output_file}')
    print(f'📄 File được lưu tại: {os.path.abspath(output_file)}')

if __name__ == '__main__':
    create_word_document()

